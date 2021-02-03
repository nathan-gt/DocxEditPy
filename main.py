# coding=utf-8
# Program created by Nathan Gueissaz-Teufel on 2021/02/02
# https://github.com/nathan-gt/DocxEditPy.git
from tkinter import Tk
from tkinter.filedialog import askopenfilename
import re
import json
import os
from docx import Document

# dictionary
# the index is the string replaced in the document, it's value is what will it be replaced with
with open('src/dictionary.json') as f:
    dictionary = json.load(f)


# function searching through the document given and replaces the string by another
# Code taken from stackoverflow user szum https://stackoverflow.com/a/42829667/13450363
def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


# User fetches the initial docx file
Tk().withdraw()
filename = askopenfilename()
doc = Document(filename)


# Loops through dictionary and file
# Code taken from stackoverflow user szum https://stackoverflow.com/a/42829667/13450363
for word, replacement in dictionary.items():
    word_re = re.compile(word)
    docx_replace_regex(doc, word_re, replacement)


# saves the file in current directory
if os.path.exists("result1.docx"):
    os.remove("result1.docx")
doc.save('result1.docx')
os.system('start result1.docx')
